"""File ingestion for SIE - local docs/spreadsheets; isolated from ChatLog.

Design goals:
- Recursively scan folder trees.
- Read ALL sheets in Excel workbooks.
- Prefer structured import when the sheet looks like society records, otherwise index as text.
- Keep society data isolated (society_id scope), with a multi-society scan mode.
"""

from __future__ import annotations

import csv
import json
import math
from datetime import date, datetime
from pathlib import Path
from typing import Any

from Backend.SIE import store

# Optional: openpyxl (xlsx), python-docx (docx), pdfplumber/pypdf (pdf), pillow+pytesseract (ocr)
_EXCEL_ERR: str | None = None
_DOCX_ERR: str | None = None
_PDF_ERR: str | None = None
_OCR_ERR: str | None = None

try:
    from openpyxl import load_workbook
except ImportError as e:
    load_workbook = None  # type: ignore[misc, assignment]
    _EXCEL_ERR = str(e)

try:
    from docx import Document as DocxDocument
except ImportError as e:
    DocxDocument = None  # type: ignore[misc, assignment]
    _DOCX_ERR = str(e)

try:
    import pdfplumber  # type: ignore[import-not-found]
except ImportError as e:
    pdfplumber = None  # type: ignore[assignment]
    _PDF_ERR = str(e)

try:
    from PIL import Image  # type: ignore[import-not-found]
    import pytesseract  # type: ignore[import-not-found]
except ImportError as e:
    Image = None  # type: ignore[assignment]
    pytesseract = None  # type: ignore[assignment]
    _OCR_ERR = str(e)


def _coerce_cell(key: str, v: Any) -> Any:
    if v is None or v == "":
        return None
    if isinstance(v, float) and (math.isnan(v) or math.isinf(v)):
        return None
    if isinstance(v, datetime):
        lk = str(key).lower()
        if "bill_month" in lk or lk == "month" or lk.endswith("_month"):
            return v.strftime("%Y-%m")
        if "date" in lk or "at" in lk or lk in ("due_date", "paid_at", "payment_date"):
            return v.strftime("%Y-%m-%d")
        return v.isoformat(sep=" ", timespec="seconds")
    if isinstance(v, date) and not isinstance(v, datetime):
        lk = str(key).lower()
        if "bill_month" in lk:
            return v.strftime("%Y-%m")
        return v.isoformat()
    if isinstance(v, float) and v == int(v):
        return int(v)
    return v


def _normalize_header(h: Any) -> str:
    s = str(h or "").strip().lower().replace(" ", "_").replace("-", "_")
    aliases = {
        "society_name": "society",
        "soc_name": "society",
        "chs_name": "society",
        "building": "society",
        "project": "society",
        "flat_no": "flat",
        "flat_number": "flat",
        "unit_no": "flat",
        "unit_number": "flat",
        "shop_no": "flat",
        "shop_number": "flat",
        "premises": "flat",
        "room_no": "flat",
        "room_number": "flat",
        "door_no": "flat",
        "door_number": "flat",
        "wing_name": "wing_name",
        "tower": "wing",
        "tower_no": "wing",
        "tower_number": "wing",
        "block": "wing",
        "bldg": "wing",
        "owner_name": "owner",
        "member_name": "owner",
        "name_of_owner": "owner",
        "tenant_name": "tenant",
        "mobile": "phone",
        "mob": "phone",
        "contact": "phone",
        "contact_no": "phone",
        "phone_no": "phone",
        "receipt_no": "receipt_no",
        "receipt_number": "receipt_no",
        "utr_no": "utr",
        "transaction_id": "utr",
        "transaction_ref": "utr",
        "cheque_number": "cheque_no",
        "check_no": "cheque_no",
        "cheque_no.": "cheque_no",
        "amount": "payment_amount",
        "paid_amount": "amount_paid",
        "due_amount": "amount_due",
        "balance": "amount_due",
        "bill_month": "bill_month",
        "billing_month": "bill_month",
        "month": "bill_month",
    }
    return aliases.get(s, s)


def normalize_spreadsheet_row(raw: dict[str, Any]) -> dict[str, Any]:
    out: dict[str, Any] = {}
    for k, v in raw.items():
        nk = _normalize_header(k)
        if not nk:
            continue
        out[nk] = _coerce_cell(nk, v)
    return out


def load_json_records(path: Path) -> list[dict[str, Any]]:
    data = json.loads(path.read_text(encoding="utf-8", errors="replace"))
    if isinstance(data, list):
        return [normalize_spreadsheet_row(x) for x in data if isinstance(x, dict)]
    if isinstance(data, dict):
        return [normalize_spreadsheet_row(data)]
    return []


def load_csv_records(path: Path) -> list[dict[str, Any]]:
    with path.open(newline="", encoding="utf-8", errors="replace") as f:
        reader = csv.DictReader(f)
        return [normalize_spreadsheet_row(dict(row)) for row in reader]


def load_xlsx_records(path: Path) -> list[dict[str, Any]]:
    if load_workbook is None:
        raise RuntimeError(
            "openpyxl is not installed. Run: pip install openpyxl"
            + (f" ({_EXCEL_ERR})" if _EXCEL_ERR else "")
        )
    wb = load_workbook(path, read_only=True, data_only=True)
    try:
        ws = wb.active
        rows = list(ws.iter_rows(values_only=True))
    finally:
        wb.close()
    if not rows:
        return []
    headers = [_normalize_header(c) for c in rows[0]]
    out: list[dict[str, Any]] = []
    for row in rows[1:]:
        d: dict[str, Any] = {}
        for i, h in enumerate(headers):
            if not h:
                continue
            val = row[i] if i < len(row) else None
            d[h] = _coerce_cell(h, val)
        if any(v is not None and str(v).strip() != "" for v in d.values()):
            out.append(d)
    return out


def load_xlsx_records_all_sheets(path: Path) -> list[dict[str, Any]]:
    """Return row dicts for ALL sheets; each row includes `_sheet` and `_source_path`."""
    if load_workbook is None:
        raise RuntimeError(
            "openpyxl is not installed. Run: pip install openpyxl"
            + (f" ({_EXCEL_ERR})" if _EXCEL_ERR else "")
        )
    wb = load_workbook(path, read_only=True, data_only=True)
    out: list[dict[str, Any]] = []
    try:
        for ws in wb.worksheets:
            rows = list(ws.iter_rows(values_only=True))
            if not rows:
                continue
            headers = [_normalize_header(c) for c in rows[0]]
            for row in rows[1:]:
                d: dict[str, Any] = {"_sheet": ws.title, "_source_path": str(path)}
                for i, h in enumerate(headers):
                    if not h:
                        continue
                    val = row[i] if i < len(row) else None
                    d[h] = _coerce_cell(h, val)
                if any(
                    (k not in {"_sheet", "_source_path"})
                    and (v is not None and str(v).strip() != "")
                    for k, v in d.items()
                ):
                    out.append(d)
    finally:
        wb.close()
    return out


def docx_to_plain_text(path: Path) -> str:
    if DocxDocument is None:
        raise RuntimeError(
            "python-docx is not installed. Run: pip install python-docx"
            + (f" ({_DOCX_ERR})" if _DOCX_ERR else "")
        )
    doc = DocxDocument(str(path))
    parts: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for tbl_row in table.rows:
            cells = [c.text.strip() for c in tbl_row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def pdf_to_plain_text(path: Path) -> str:
    if pdfplumber is None:
        raise RuntimeError(
            "pdfplumber is not installed. Run: pip install pdfplumber"
            + (f" ({_PDF_ERR})" if _PDF_ERR else "")
        )
    parts: list[str] = []
    with pdfplumber.open(str(path)) as pdf:
        for i, page in enumerate(pdf.pages):
            try:
                text = page.extract_text() or ""
            except Exception:
                text = ""
            text = text.strip()
            if text:
                parts.append(f"[page {i+1}]\n{text}")
    return "\n\n".join(parts)


def image_to_text(path: Path) -> str:
    if Image is None or pytesseract is None:
        raise RuntimeError(
            "OCR dependencies not installed. Run: pip install pillow pytesseract"
            + (f" ({_OCR_ERR})" if _OCR_ERR else "")
        )
    img = Image.open(str(path))
    try:
        return (pytesseract.image_to_string(img) or "").strip()
    finally:
        try:
            img.close()
        except Exception:
            pass


def _looks_like_structured_society_rows(recs: list[dict[str, Any]]) -> bool:
    """Heuristic: do we have enough signals to try structured import?"""
    if not recs:
        return False
    score = 0
    keys = set()
    for r in recs[:50]:
        keys.update(k for k, v in r.items() if v is not None)
    low = {str(k).lower() for k in keys}
    if "society" in low or "society_name" in low:
        score += 3
    if "flat" in low or "unit" in low or "unit_label" in low:
        score += 2
    if "wing" in low or "wing_code" in low or "tower" in low or "block" in low:
        score += 1
    if "owner" in low or "tenant" in low:
        score += 1
    if "bill_month" in low or "amount_due" in low or "amount_paid" in low:
        score += 1
    if "utr" in low or "cheque_no" in low or "payment_amount" in low:
        score += 1
    return score >= 4


def _ensure_society(conn, name: str) -> int:
    nm = (name or "").strip()
    if not nm:
        raise ValueError("empty society name")
    row = conn.execute("SELECT id FROM societies WHERE name = ?", (nm,)).fetchone()
    if row:
        return int(row["id"])
    conn.execute("INSERT INTO societies (name, aliases_json) VALUES (?, '[]')", (nm,))
    return int(conn.execute("SELECT last_insert_rowid()").fetchone()[0])


def scan_folder(
    root: Path,
    *,
    society_id: int | None = None,
    import_rows: bool = True,
) -> str:
    root = root.resolve()
    if not root.is_dir():
        return f"Not a directory: {root}"

    imported_docs = 0
    rows_imported = 0
    pay_n = 0
    warnings: list[str] = []

    with store.get_connection() as conn:
        # Multi-society mode: if scanning a parent folder, treat each first-level folder as a society bucket.
        if society_id is None:
            top_dirs = [d for d in root.iterdir() if d.is_dir()]
            root_files = [f for f in root.iterdir() if f.is_file()]
            # If it's mostly subfolders, assume each is a society unless proven otherwise.
            if len(top_dirs) >= 2 and len(root_files) <= 5:
                for d in top_dirs:
                    sid = _ensure_society(conn, d.name)
                    # recursive scan per society folder
                    msg = scan_folder(d, society_id=sid, import_rows=import_rows)
                    warnings.append(msg)
                store.audit(conn, "scan_folder_multi", str(root))
                return "Multi-society scan complete.\n" + "\n".join(warnings[:25])

            # Single-society: default society name to folder name.
            society_id = _ensure_society(conn, root.name)

        for p in root.rglob("*"):
            if not p.is_file():
                continue
            suf = p.suffix.lower()
            try:
                if suf in {".txt", ".md"}:
                    body = p.read_text(encoding="utf-8", errors="replace")
                    store.insert_document(
                        conn,
                        society_id=society_id,
                        source_path=str(p),
                        title=p.name,
                        doc_type="text",
                        body_text=body,
                    )
                    imported_docs += 1
                elif suf == ".json" and import_rows:
                    recs = load_json_records(p)
                    is_structured = _looks_like_structured_society_rows(recs)
                    if is_structured:
                        # If sheet doesn't contain society column, pin to scan scope society_id.
                        for r in recs:
                            if "society" not in r and "society_name" not in r and society_id is not None:
                                r["society"] = conn.execute(
                                    "SELECT name FROM societies WHERE id = ?", (society_id,)
                                ).fetchone()["name"]
                        f, pay = store.import_record_bundle(conn, recs)
                        rows_imported += f
                        pay_n += pay
                    else:
                        body = p.read_text(encoding="utf-8", errors="replace")
                        store.insert_document(
                            conn,
                            society_id=society_id,
                            source_path=str(p),
                            title=p.name,
                            doc_type="json",
                            body_text=body,
                        )
                        imported_docs += 1
                elif suf == ".csv" and import_rows:
                    recs = load_csv_records(p)
                    if _looks_like_structured_society_rows(recs):
                        for r in recs:
                            if "society" not in r and "society_name" not in r and society_id is not None:
                                r["society"] = conn.execute(
                                    "SELECT name FROM societies WHERE id = ?", (society_id,)
                                ).fetchone()["name"]
                    f, pay = store.import_record_bundle(conn, recs)
                    rows_imported += f
                    pay_n += pay
                elif suf in {".xlsx", ".xlsm"} and import_rows:
                    try:
                        recs = load_xlsx_records_all_sheets(p)
                    except RuntimeError as e:
                        warnings.append(str(e))
                        continue
                    is_structured = _looks_like_structured_society_rows(recs)
                    if is_structured:
                        for r in recs:
                            if "society" not in r and "society_name" not in r and society_id is not None:
                                r["society"] = conn.execute(
                                    "SELECT name FROM societies WHERE id = ?", (society_id,)
                                ).fetchone()["name"]
                        f, pay = store.import_record_bundle(conn, recs)
                        rows_imported += f
                        pay_n += pay
                    else:
                        # Index ALL sheets as text, grouped by sheet.
                        by_sheet: dict[str, list[dict[str, Any]]] = {}
                        for r in recs:
                            by_sheet.setdefault(str(r.get("_sheet") or "Sheet"), []).append(r)
                        parts: list[str] = []
                        for sh, rows in list(by_sheet.items())[:50]:
                            parts.append(f"[sheet] {sh}")
                            for r in rows[:200]:
                                parts.append(
                                    " ; ".join(
                                        f"{k}={v}"
                                        for k, v in sorted(r.items())
                                        if k not in {"_sheet", "_source_path"} and v is not None
                                    )
                                )
                        body = "\n".join([p for p in parts if p.strip()])
                        if not body.strip():
                            body = "(empty or unreadable sheet)"
                        store.insert_document(
                            conn,
                            society_id=society_id,
                            source_path=str(p),
                            title=p.name,
                            doc_type="excel",
                            body_text=body[:500000],
                        )
                        imported_docs += 1
                elif suf == ".docx":
                    try:
                        body = docx_to_plain_text(p)
                    except RuntimeError as e:
                        warnings.append(str(e))
                        continue
                    store.insert_document(
                        conn,
                        society_id=society_id,
                        source_path=str(p),
                        title=p.stem,
                        doc_type="word",
                        body_text=body,
                    )
                    imported_docs += 1
                elif suf == ".pdf":
                    try:
                        body = pdf_to_plain_text(p)
                    except RuntimeError as e:
                        warnings.append(str(e))
                        continue
                    if not body.strip():
                        body = "(no extractable text; try OCR on exported images if needed)"
                    store.insert_document(
                        conn,
                        society_id=society_id,
                        source_path=str(p),
                        title=p.name,
                        doc_type="pdf",
                        body_text=body[:800000],
                    )
                    imported_docs += 1
                elif suf in {".png", ".jpg", ".jpeg", ".webp", ".tiff", ".bmp"}:
                    try:
                        body = image_to_text(p)
                    except RuntimeError as e:
                        warnings.append(str(e))
                        continue
                    if body.strip():
                        store.insert_document(
                            conn,
                            society_id=society_id,
                            source_path=str(p),
                            title=p.name,
                            doc_type="image_ocr",
                            body_text=body[:200000],
                        )
                        imported_docs += 1
            except OSError:
                continue
            except Exception as e:
                warnings.append(f"{p.name}: {e}")
                continue
        store.audit(conn, "scan_folder", str(root))

    msg = (
        f"Scan complete: {imported_docs} documents indexed, "
        f"{rows_imported} flat links touched, {pay_n} new payments (duplicates skipped)."
    )
    if warnings:
        msg += " Warnings: " + "; ".join(warnings[:5])
        if len(warnings) > 5:
            msg += f" (+{len(warnings) - 5} more)"
    return msg
